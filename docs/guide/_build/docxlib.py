"""
docxlib.py — a small, reusable Word (.docx) builder on top of python-docx.

Why this exists
---------------
The CryptoCopilot guide is produced from a single set of helper functions, so the
styling stays identical across the whole document and the build stays simple.

Everything here is deterministic and reproducible: `python3 build_guides.py`
regenerates the .docx file from scratch.
"""
from __future__ import annotations

import os
from PIL import Image

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ----------------------------------------------------------------------------
# Palette (a calm fintech navy + teal/blue accents, matching the dark app UI)
# ----------------------------------------------------------------------------
NAVY    = RGBColor(0x0F, 0x2B, 0x46)   # deep navy — H1
BLUE    = RGBColor(0x1D, 0x4E, 0x89)   # primary blue — H2
TEAL    = RGBColor(0x0E, 0x8F, 0x7E)   # accent teal — H3 / highlights
SLATE   = RGBColor(0x33, 0x3A, 0x45)   # body text
GREY    = RGBColor(0x6B, 0x72, 0x80)   # captions / meta
LIGHTBG = "EAF1F8"                       # callout background (hex string, no #)
CODEBG  = "F2F3F5"                       # code background
HEADBG  = "12243A"                       # table header background (dark navy)
ZEBRA   = "F4F7FB"                       # table zebra row
RULE    = "C9D6E5"                       # table border colour

LATIN_FONT   = "Calibri"
MONO_FONT    = "Consolas"

# Friendly aliases -> screenshot files (resolved at build time)
SHOTS_DIR = os.path.join(os.path.dirname(__file__), "..", "screenshots")
SHOTS = {
    "signals":     "Screenshot 2026-06-06 at 11.39.19.png",
    "markets":     "Screenshot 2026-06-06 at 11.39.54.png",
    "ml_pipeline": "Screenshot 2026-06-06 at 11.40.12.png",
    "coin_detail": "Screenshot 2026-06-06 at 11.41.01.png",
    "signal_pair": "Screenshot 2026-06-06 at 11.41.41.png",
    "signal_card": "Screenshot 2026-06-06 at 11.42.10.png",
    "analyst":     "Screenshot 2026-06-06 at 11.42.56.png",
    "paper_trades":"Screenshot 2026-06-06 at 11.43.40.png",
    "performance": "Screenshot 2026-06-06 at 11.43.53.png",
    "er_diagram":  "Screenshot 2026-06-06 at 12.02.07.png",
}

def shot(alias: str) -> str:
    return os.path.normpath(os.path.join(SHOTS_DIR, SHOTS[alias]))


# ----------------------------------------------------------------------------
# Low-level OOXML helpers
# ----------------------------------------------------------------------------
def _set(el, tag, **attrs):
    child = OxmlElement(tag)
    for k, v in attrs.items():
        child.set(qn(k), v)
    el.append(child)
    return child

def _shade(el_pr, fill):
    _set(el_pr, "w:shd", **{"w:val": "clear", "w:color": "auto", "w:fill": fill})

def _run_fonts(run, latin, cs, size_pt=None):
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts"); rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), latin)
    rFonts.set(qn("w:hAnsi"), latin)
    rFonts.set(qn("w:cs"), cs)
    if size_pt is not None:
        szCs = OxmlElement("w:szCs"); szCs.set(qn("w:val"), str(int(size_pt * 2)))
        rPr.append(szCs)


# ----------------------------------------------------------------------------
# Builder
# ----------------------------------------------------------------------------
class Doc:
    def __init__(self):
        self.latin = LATIN_FONT
        self.cs = LATIN_FONT
        self.d = Document()
        self._setup_styles()
        self._setup_page()

    # -- document setup -----------------------------------------------------
    def _setup_styles(self):
        d = self.d
        normal = d.styles["Normal"]
        normal.font.name = self.latin
        normal.font.size = Pt(10.5)
        normal.font.color.rgb = SLATE
        rpr = normal.element.get_or_add_rPr()
        rFonts = rpr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rpr.insert(0, rFonts)
        rFonts.set(qn("w:ascii"), self.latin)
        rFonts.set(qn("w:hAnsi"), self.latin)
        rFonts.set(qn("w:cs"), self.cs)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.15

    def _setup_page(self):
        for s in self.d.sections:
            s.top_margin = Inches(0.85)
            s.bottom_margin = Inches(0.85)
            s.left_margin = Inches(0.9)
            s.right_margin = Inches(0.9)

    def _make_update_fields(self):
        # ask Word to refresh fields (the TOC) when the document is opened
        settings = self.d.settings.element
        if settings.find(qn("w:updateFields")) is None:
            uf = OxmlElement("w:updateFields"); uf.set(qn("w:val"), "true")
            settings.append(uf)

    def enable_update_fields(self):
        self._make_update_fields()

    # -- primitives ---------------------------------------------------------
    def _style_run(self, run, size, color, bold=False, italic=False, mono=False):
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run.bold = bold
        run.italic = italic
        latin = MONO_FONT if mono else self.latin
        cs = MONO_FONT if mono else self.cs
        _run_fonts(run, latin, cs, size_pt=size)

    def _emit_rich(self, p, text, size, color, base_bold=False):
        """Render text with **bold** and `code` inline spans."""
        import re
        # tokenise on **bold** and `code`
        parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
        for part in parts:
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                r = p.add_run(part[2:-2]); self._style_run(r, size, color, bold=True)
            elif part.startswith("`") and part.endswith("`"):
                r = p.add_run(part[1:-1]); self._style_run(r, size, BLUE, mono=True)
            else:
                r = p.add_run(part); self._style_run(r, size, color, bold=base_bold)

    # -- public blocks ------------------------------------------------------
    def spacer(self, pts=6):
        p = self.d.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(""); r.font.size = Pt(pts)
        return p

    def page_break(self):
        self.d.add_page_break()

    def h1(self, text, num=None):
        p = self.d.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        p.style = self.d.styles["Heading 1"]
        label = f"{num}.  {text}" if num else text
        r = p.add_run(label)
        self._style_run(r, 19, NAVY, bold=True)
        self._bottom_border(p, "1D4E89", 12)
        return p

    def h2(self, text):
        p = self.d.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        p.style = self.d.styles["Heading 2"]
        r = p.add_run(text); self._style_run(r, 14.5, BLUE, bold=True)
        return p

    def h3(self, text):
        p = self.d.add_paragraph()
        p.paragraph_format.space_before = Pt(9)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.keep_with_next = True
        p.style = self.d.styles["Heading 3"]
        r = p.add_run(text); self._style_run(r, 12, TEAL, bold=True)
        return p

    def p(self, text, size=10.5, color=None, align=None):
        para = self.d.add_paragraph()
        self._emit_rich(para, text, size, color or SLATE)
        if align is not None:
            para.alignment = align
        return para

    def bullet(self, text, size=10.5, level=0):
        para = self.d.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        para.paragraph_format.space_after = Pt(3)
        self._emit_rich(para, text, size, SLATE)
        return para

    def numbered(self, text, size=10.5):
        para = self.d.add_paragraph(style="List Number")
        para.paragraph_format.space_after = Pt(3)
        self._emit_rich(para, text, size, SLATE)
        return para

    def callout(self, text, fill=LIGHTBG, accent="1D4E89", size=10):
        """A shaded note box with a left accent border."""
        p = self.d.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.left_indent = Inches(0.08)
        p.paragraph_format.right_indent = Inches(0.08)
        pPr = p._p.get_or_add_pPr()
        _shade(pPr, fill)
        # border box
        pbdr = OxmlElement("w:pBdr")
        for side in ("top", "left", "bottom", "right"):
            b = OxmlElement(f"w:{side}")
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), "18" if side == "left" else "4")
            b.set(qn("w:space"), "6")
            b.set(qn("w:color"), accent if side == "left" else RULE)
            pbdr.append(b)
        pPr.append(pbdr)
        self._emit_rich(p, text, size, SLATE)
        return p

    def note(self, text, size=9.5, color=None):
        """A small italic muted line (e.g. an on-screen cue in the script)."""
        p = self.d.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text); r.italic = True
        r.font.size = Pt(size); r.font.color.rgb = color or GREY
        _run_fonts(r, self.latin, self.cs, size_pt=size)
        return p

    def code(self, text, size=8.8):
        """Monospace, shaded code block."""
        lines = text.strip("\n").split("\n")
        p = self.d.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pPr = p._p.get_or_add_pPr()
        _shade(pPr, CODEBG)
        pbdr = OxmlElement("w:pBdr")
        for side in ("top", "left", "bottom", "right"):
            b = OxmlElement(f"w:{side}")
            b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "4")
            b.set(qn("w:space"), "4"); b.set(qn("w:color"), "DCE0E6")
            pbdr.append(b)
        pPr.append(pbdr)
        for i, line in enumerate(lines):
            r = p.add_run(line)
            r.font.name = MONO_FONT; r.font.size = Pt(size)
            r.font.color.rgb = RGBColor(0x24, 0x2A, 0x33)
            _run_fonts(r, MONO_FONT, MONO_FONT, size_pt=size)
            if i < len(lines) - 1:
                r.add_break()
        return p

    def table(self, headers, rows, widths=None, header_fill=HEADBG, font=9.0):
        n = len(headers)
        t = self.d.add_table(rows=1, cols=n)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = True
        self._table_borders(t)
        # header
        hdr = t.rows[0].cells
        for j, htext in enumerate(headers):
            cell = hdr[j]
            self._fill_cell(cell, header_fill)
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.space_before = Pt(2)
            r = para.add_run(htext)
            r.font.bold = True; r.font.size = Pt(font + 0.3)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _run_fonts(r, self.latin, self.cs, size_pt=font + 0.3)
        # body
        for i, row in enumerate(rows):
            cells = t.add_row().cells
            for j, val in enumerate(row):
                cell = cells[j]
                if i % 2 == 1:
                    self._fill_cell(cell, ZEBRA)
                para = cell.paragraphs[0]
                para.paragraph_format.space_after = Pt(2)
                para.paragraph_format.space_before = Pt(2)
                self._emit_rich(para, str(val), font, SLATE)
        if widths:
            # fixed layout so Word honours the column widths exactly (no overflow)
            t.autofit = False
            tblPr = t._tbl.tblPr
            layout = OxmlElement("w:tblLayout"); layout.set(qn("w:type"), "fixed")
            tblPr.append(layout)
            tblW = OxmlElement("w:tblW")
            tblW.set(qn("w:type"), "dxa"); tblW.set(qn("w:w"), str(int(sum(widths) * 1440)))
            tblPr.append(tblW)
            for row in t.rows:
                for j, w in enumerate(widths):
                    row.cells[j].width = Inches(w)
        self.spacer(2)
        return t

    def image(self, alias_or_path, caption=None, max_w=6.4, max_h=7.2):
        path = shot(alias_or_path) if alias_or_path in SHOTS else alias_or_path
        w_px, h_px = Image.open(path).size
        ratio = w_px / h_px
        w_in = max_w
        h_in = w_in / ratio
        if h_in > max_h:
            h_in = max_h
            w_in = h_in * ratio
        p = self.d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run()
        run.add_picture(path, width=Inches(w_in))
        # thin frame around the figure
        self._bottom_border(p, "C9D6E5", 4)
        if caption:
            cap = self.d.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_after = Pt(10)
            r = cap.add_run(caption)
            r.italic = True; r.font.size = Pt(8.8); r.font.color.rgb = GREY
            _run_fonts(r, self.latin, self.cs, size_pt=8.8)
        return p

    def toc(self, title):
        """Insert a real, updatable Word table-of-contents field."""
        self._make_update_fields()
        p = self.d.add_paragraph()
        fld_begin = OxmlElement("w:fldSimple")
        fld_begin.set(qn("w:instr"), r'TOC \o "1-3" \h \z \u')
        run_holder = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = "Right-click here and choose “Update Field” to build the contents."
        rpr = OxmlElement("w:rPr")
        it = OxmlElement("w:i"); rpr.append(it)
        run_holder.append(rpr); run_holder.append(t)
        fld_begin.append(run_holder)
        p._p.append(fld_begin)
        return p

    # -- helpers for tables / borders --------------------------------------
    def _fill_cell(self, cell, fill):
        tcPr = cell._tc.get_or_add_tcPr()
        _shade(tcPr, fill)

    def _table_borders(self, t):
        tbl = t._tbl
        tblPr = tbl.tblPr
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            e = OxmlElement(f"w:{edge}")
            e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "4")
            e.set(qn("w:space"), "0"); e.set(qn("w:color"), RULE)
            borders.append(e)
        tblPr.append(borders)

    def _bottom_border(self, p, color, sz):
        pPr = p._p.get_or_add_pPr()
        pbdr = pPr.find(qn("w:pBdr"))
        if pbdr is None:
            pbdr = OxmlElement("w:pBdr"); pPr.append(pbdr)
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), str(sz))
        bottom.set(qn("w:space"), "3"); bottom.set(qn("w:color"), color)
        pbdr.append(bottom)

    # -- title page ---------------------------------------------------------
    def title_page(self, title, subtitle, tagline, meta_lines):
        self.spacer(70)
        p = self.d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title); r.bold = True; r.font.size = Pt(34); r.font.color.rgb = NAVY
        _run_fonts(r, self.latin, self.cs, 34)
        p2 = self.d.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p2.add_run(subtitle); r.font.size = Pt(15); r.font.color.rgb = TEAL; r.bold = True
        _run_fonts(r, self.latin, self.cs, 15)
        self.spacer(6)
        p3 = self.d.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p3.add_run(tagline); r.italic = True; r.font.size = Pt(11.5); r.font.color.rgb = GREY
        _run_fonts(r, self.latin, self.cs, 11.5)
        self.spacer(30)
        for line in meta_lines:
            pp = self.d.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = pp.add_run(line); r.font.size = Pt(10.5); r.font.color.rgb = SLATE
            _run_fonts(r, self.latin, self.cs, 10.5)
        self.page_break()

    def save(self, path):
        os.makedirs(os.path.dirname(path), exist_ok=True)
        self.d.save(path)
        return path
